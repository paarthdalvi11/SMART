from fastapi import FastAPI, Request
from fastapi.responses import FileResponse
from docx import Document
import tempfile

app = FastAPI()

@app.post("/api/export")
async def export_docx(request: Request):
    reqs = await request.json()
    doc = Document()
    doc.add_heading("Extracted Requirements", 0)
    for req in reqs:
        doc.add_heading(f"{req['id']} ({req['type']})", level=2)
        doc.add_paragraph(f"Description: {req['description']}")
        doc.add_paragraph(f"Priority: {req['priority']}")
        if req.get("note"):
            doc.add_paragraph(f"Note: {req['note']}")
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        doc.save(tmp.name)
        return FileResponse(tmp.name, filename="requirements.docx")