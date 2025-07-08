import os
import json
import sys
from docx import Document
from dotenv import load_dotenv
import google.generativeai as genai

# === Load API Key ===
load_dotenv()
genai.configure(api_key=os.getenv("GEMINI_API"))

# === Consolidate Requirements Function ===
def consolidate_requirements(json_dicts):
    combined = []
    for source, req_list in json_dicts.items():
        combined.append(f"From {source}:\n{json.dumps(req_list, indent=2)}\n")
    
    print("COMBINED:", combined, file=sys.stderr)
    combined_text = "\n\n".join(combined)

    print("COMBINED TEXT:", combined_text, file=sys.stderr)
    prompt = f"""
You are a software requirements engineer. The following are multiple lists of system requirements (some might be overlapping or redundant) extracted from various inputs.

Task:
1. Deduplicate overlapping or similar requirements.
2. Prioritize using MoSCoW (must/should/could/won't).
3. Classify each requirement into:
   - 'functional' (system behavior, features)
   - 'non-functional' (performance, security, usability, etc.)

Return the result as a structured JSON object with two main keys:
- "functional_requirements": a list of deduplicated functional requirements.
- "non_functional_requirements": a list of deduplicated non-functional requirements.

Each item should be an object with:
- "id" (e.g., FR-1, NFR-1)
- "description"
- "priority" (must, should, could, won't)
- "type" (functional, non-functional)
- "note" (optional note for the requirement)

Input:
{combined_text}
"""

    # Call Gemini API for processing
    # model = genai.GenerativeModel("gemini-1.5-pro-latest")
    model = genai.GenerativeModel(os.getenv("GEMINI_MODEL"))
    response = model.generate_content(prompt)
    result_text = response.text
    
    print("\n\nRESULT TEXT:", result_text, file=sys.stderr)

    try:
        # Extract JSON from the response
        json_start = result_text.find("{")
        json_end = result_text.rfind("}") + 1
        return json.loads(result_text[json_start:json_end])
    except Exception as e:
        print("Failed to parse JSON from LLM response:\n", result_text, file=sys.stderr)
        raise e

# === Write to DOCX Function ===
def write_to_docx(requirements, output_file="Final_Requirements.docx"):
    doc = Document()
    doc.add_heading("Final Consolidated Requirements", 0)

    for section, title in [("functional_requirements", "Functional Requirements"), ("non_functional_requirements", "Non-Functional Requirements")]:
        doc.add_heading(title, level=1)
        for req in requirements.get(section, []):
            para = doc.add_paragraph()
            para.add_run(f"{req['id']}: ").bold = True
            para.add_run(req['description'])
            para.add_run(f"\nPriority: {req.get('priority', 'N/A')}")
            if 'note' in req:
                para.add_run(f"\nNote: {req['note']}")

    doc.save(output_file)
    print(f"DOCX created: {output_file}", file=sys.stderr)
    return output_file

# === Orchestrate Final Workflow ===
def orchestrate_final_processing(all_jsons, output_path):
    print("Starting final processing...", file=sys.stderr)

    # Consolidate requirements from all JSONs
    consolidated_requirements = consolidate_requirements(all_jsons)

    # Write the consolidated requirements to a DOCX file
    write_to_docx(consolidated_requirements, output_file=output_path)

    print("\nFinal processing complete.", file=sys.stderr)

# Example: Assuming `all_jsons`